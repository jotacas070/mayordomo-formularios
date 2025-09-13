import streamlit as st
import requests
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
import base64

# Configuración de la página
st.set_page_config(
    page_title="Mayordomo General IA - Formularios",
    page_icon="⚓",
    layout="wide"
)

# URL de tu API de Flowise
FLOWISE_API_URL = "https://flowiseai-railway-production-eb75.up.railway.app/api/v1/prediction/0b96fddb-9b55-48d2-8d22-5e92b0710781"

# Definición de formularios
FORMULARIOS = {
    "sancionamiento_material": {
        "nombre": "Formulario de Sancionamiento de Material",
        "tipo": "Certificación",
        "departamento": "Centro de Abastecimiento",
        "campos": [
            {"id": "centro_abastecimiento", "label": "Centro de Abastecimiento", "tipo": "text", "requerido": True},
            {"id": "asesor_tec_nombre", "label": "Nombre del Asesor Técnico", "tipo": "text", "requerido": True},
            {"id": "uurr", "label": "UU.RR. (Unidad/Repartición)", "tipo": "text", "requerido": True},
            {"id": "fecha_sancionamiento", "label": "Fecha de sancionamiento", "tipo": "date", "requerido": True},
            {"id": "tipo_clasificacion", "label": "Tipo de clasificación", "tipo": "select", "opciones": ["Para depósito", "Para destrucción", "Mixto"], "requerido": True},
            {"id": "lugar", "label": "Lugar", "tipo": "text", "requerido": True},
            {"id": "fecha_emision", "label": "Fecha de emisión", "tipo": "date", "requerido": True},
            {"id": "nombre_jefe_centro", "label": "Nombre Jefe Centro AB.", "tipo": "text", "requerido": True},
            {"id": "reparticion", "label": "Repartición", "tipo": "text", "requerido": True}
        ],
        "items_list": {"id": "material_items", "label": "Materiales", "campos": ["NUS/NAS", "N°Serie", "Cantidad", "Descripción", "Valor Contable"]}
    },
    "guia_traspaso": {
        "nombre": "Guía de Traspaso de Material Inventariable",
        "tipo": "Guía de Traspaso", 
        "departamento": "Abastecimiento",
        "campos": [
            {"id": "reparticion_recibe", "label": "Repartición que recibe", "tipo": "text", "requerido": True},
            {"id": "documento_autorizacion", "label": "Documento de autorización", "tipo": "text", "requerido": True},
            {"id": "nombre_comandante_envia", "label": "Comandante que envía", "tipo": "text", "requerido": True},
            {"id": "uurr_actual", "label": "UU.RR. Actual (Código)", "tipo": "text", "requerido": True},
            {"id": "nombre_uurr_futura", "label": "UU.RR. Futura", "tipo": "text", "requerido": True},
            {"id": "lugar_fecha", "label": "Lugar y fecha", "tipo": "text", "requerido": True}
        ],
        "items_list": {"id": "items_traspaso", "label": "Items a traspasar", "campos": ["NAS/NUS", "Cantidad", "Descripción", "Sección"]}
    },
    "acta_perdida": {
        "nombre": "Acta de Constatación de Pérdida",
        "tipo": "Acta",
        "departamento": "Unidad Operativa", 
        "campos": [
            {"id": "uurr", "label": "UU.RR. (Unidad/Repartición)", "tipo": "text", "requerido": True},
            {"id": "motivo_perdida", "label": "Motivo de la Pérdida", "tipo": "textarea", "requerido": True},
            {"id": "valor_total_utm", "label": "Valor vs 6 UTM", "tipo": "select", "opciones": ["Inferior a 6 UTM", "Superior a 6 UTM"], "requerido": True},
            {"id": "lugar", "label": "Lugar", "tipo": "text", "requerido": True},
            {"id": "fecha", "label": "Fecha", "tipo": "date", "requerido": True},
            {"id": "nombre_asesor_tecnico", "label": "Asesor Técnico", "tipo": "text", "requerido": True},
            {"id": "nombre_jefe_centro", "label": "Jefe Centro AB.", "tipo": "text", "requerido": True}
        ],
        "items_list": {"id": "items_perdida", "label": "Items Perdidos", "campos": ["NUS/NAS", "Cantidad", "Descripción", "Valor Contable"]}
    }
}

def inicializar_session_state():
    """Inicializa las variables de sesión"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'formulario_activo' not in st.session_state:
        st.session_state.formulario_activo = None
    if 'datos_formulario' not in st.session_state:
        st.session_state.datos_formulario = {}
    if 'items_formulario' not in st.session_state:
        st.session_state.items_formulario = []
    if 'paso_actual' not in st.session_state:
        st.session_state.paso_actual = 0

def consultar_flowise(mensaje):
    """Consulta la API de Flowise"""
    try:
        payload = {
            "question": mensaje,
            "overrideConfig": {}
        }
        
        response = requests.post(
            FLOWISE_API_URL,
            json=payload,
            headers={"Content-Type": "application/json"}
        )
        
        if response.status_code == 200:
            return response.json().get("text", "Error al procesar respuesta")
        else:
            return f"Error al consultar el sistema: {response.status_code}"
            
    except Exception as e:
        return f"Error de conexión: {str(e)}"

def detectar_formulario_necesario(mensaje):
    """Detecta si el usuario necesita un formulario específico"""
    mensaje_lower = mensaje.lower()
    
    # Keywords para detección
    formularios_keywords = {
        "sancionamiento_material": ["sancionar", "sancionamiento", "clasificar", "destrucción", "depósito", "baja"],
        "guia_traspaso": ["traspaso", "traspasar", "transferir", "enviar", "material", "unidad"],
        "acta_perdida": ["pérdida", "perdido", "constatación", "faltante", "extraviado"]
    }
    
    for form_id, keywords in formularios_keywords.items():
        if any(keyword in mensaje_lower for keyword in keywords):
            return form_id
    
    return None

def renderizar_campo(campo, valor_actual=None):
    """Renderiza un campo del formulario según su tipo"""
    campo_id = campo["id"]
    label = campo["label"]
    
    if campo["tipo"] == "text":
        return st.text_input(label, value=valor_actual or "", key=campo_id)
    elif campo["tipo"] == "textarea":
        return st.text_area(label, value=valor_actual or "", key=campo_id)
    elif campo["tipo"] == "date":
        return st.date_input(label, key=campo_id)
    elif campo["tipo"] == "select":
        opciones = campo.get("opciones", [])
        index = 0
        if valor_actual and valor_actual in opciones:
            index = opciones.index(valor_actual)
        return st.selectbox(label, opciones, index=index, key=campo_id)
    else:
        return st.text_input(label, value=valor_actual or "", key=campo_id)

def generar_documento_word(form_id, datos):
    """Genera documento Word con los datos del formulario"""
    doc = Document()
    formulario = FORMULARIOS[form_id]
    
    # Título
    title = doc.add_heading(formulario["nombre"], 0)
    title.alignment = 1  # Centro
    
    # Información del documento
    doc.add_paragraph(f"Tipo: {formulario['tipo']}")
    doc.add_paragraph(f"Departamento: {formulario['departamento']}")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Campos del formulario
    doc.add_heading("INFORMACIÓN DEL FORMULARIO:", level=1)
    
    for campo in formulario["campos"]:
        if campo["id"] in datos:
            valor = datos[campo["id"]]
            if valor:  # Solo agregar si tiene valor
                p = doc.add_paragraph()
                p.add_run(f"{campo['label']}: ").bold = True
                p.add_run(str(valor))
    
    # Items si existen
    if "items_list" in formulario and st.session_state.items_formulario:
        doc.add_paragraph("")
        doc.add_heading(f"{formulario['items_list']['label']}:", level=1)
        
        # Crear tabla
        table = doc.add_table(rows=1, cols=len(formulario["items_list"]["campos"]))
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(formulario["items_list"]["campos"]):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Datos
        for item in st.session_state.items_formulario:
            row_cells = table.add_row().cells
            for i, valor in enumerate(item.values()):
                if i < len(row_cells):
                    row_cells[i].text = str(valor)
    
    # Firmas
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_heading("FIRMAS Y AUTORIZACIONES:", level=1)
    
    firmas_table = doc.add_table(rows=3, cols=2)
    firmas_table.style = 'Table Grid'
    
    firmas_table.cell(0, 0).text = "SOLICITANTE:"
    firmas_table.cell(0, 1).text = "_" * 40
    firmas_table.cell(1, 0).text = "AUTORIZA:"  
    firmas_table.cell(1, 1).text = "_" * 40
    firmas_table.cell(2, 0).text = "FECHA:"
    firmas_table.cell(2, 1).text = "_" * 40
    
    return doc

def main():
    st.title("⚓ Mayordomo General IA - Gestión de Formularios")
    st.markdown("**Sistema Inteligente de Formularios Oficiales de la Armada**")
    
    inicializar_session_state()
    
    # Sidebar con información
    with st.sidebar:
        st.header("📋 Formularios Disponibles")
        for form_id, form_data in FORMULARIOS.items():
            st.write(f"**{form_data['nombre']}**")
            st.write(f"_{form_data['tipo']} - {form_data['departamento']}_")
            st.write("")
        
        if st.session_state.formulario_activo:
            st.success(f"Formulario activo: {FORMULARIOS[st.session_state.formulario_activo]['nombre']}")
            if st.button("Cancelar formulario"):
                st.session_state.formulario_activo = None
                st.session_state.datos_formulario = {}
                st.session_state.items_formulario = []
                st.session_state.paso_actual = 0
                st.rerun()
    
    # Chat interface o formulario
    if not st.session_state.formulario_activo:
        # Modo chat normal
        st.header("💬 Asistente Virtual")
        
        # Mostrar mensajes
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Input del usuario
        if prompt := st.chat_input("Escribe tu consulta..."):
            # Agregar mensaje del usuario
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Detectar si necesita formulario
            form_needed = detectar_formulario_necesario(prompt)
            
            if form_needed and form_needed in FORMULARIOS:
                st.session_state.formulario_activo = form_needed
                st.session_state.datos_formulario = {}
                st.session_state.items_formulario = []
                st.session_state.paso_actual = 0
                
                response = f"He detectado que necesitas completar: **{FORMULARIOS[form_needed]['nombre']}**. Te guiaré paso a paso para completarlo."
                
                with st.chat_message("assistant"):
                    st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
                st.rerun()
                
            else:
                # Consultar Flowise para respuesta normal
                with st.chat_message("assistant"):
                    with st.spinner("Consultando..."):
                        response = consultar_flowise(prompt)
                        st.markdown(response)
                        
                        # Verificar si la respuesta sugiere un formulario
                        form_suggested = detectar_formulario_necesario(response)
                        if form_suggested:
                            st.info("💡 ¿Necesitas completar algún formulario? Dime 'sí' y te ayudo.")
                
                st.session_state.messages.append({"role": "assistant", "content": response})
    
    else:
        # Modo formulario activo
        formulario = FORMULARIOS[st.session_state.formulario_activo]
        
        st.header(f"📝 {formulario['nombre']}")
        st.write(f"**{formulario['tipo']}** - {formulario['departamento']}")
        
        # Crear tabs
        tab1, tab2, tab3 = st.tabs(["📋 Datos Principales", "📦 Items/Materiales", "📄 Generar Documento"])
        
        with tab1:
            st.subheader("Información Principal")
            
            # Renderizar campos
            for campo in formulario["campos"]:
                valor_actual = st.session_state.datos_formulario.get(campo["id"])
                valor = renderizar_campo(campo, valor_actual)
                st.session_state.datos_formulario[campo["id"]] = valor
        
        with tab2:
            if "items_list" in formulario:
                st.subheader(formulario["items_list"]["label"])
                
                # Mostrar items existentes
                if st.session_state.items_formulario:
                    st.write("**Items agregados:**")
                    for i, item in enumerate(st.session_state.items_formulario):
                        cols = st.columns([3, 1])
                        with cols[0]:
                            item_display = " | ".join([f"{k}: {v}" for k, v in item.items() if v])
                            st.write(f"{i+1}. {item_display}")
                        with cols[1]:
                            if st.button("Eliminar", key=f"del_{i}"):
                                st.session_state.items_formulario.pop(i)
                                st.rerun()
                
                # Agregar nuevo item
                st.write("**Agregar nuevo item:**")
                campos_item = formulario["items_list"]["campos"]
                
                nuevo_item = {}
                cols = st.columns(len(campos_item))
                
                for i, campo in enumerate(campos_item):
                    with cols[i]:
                        valor = st.text_input(campo, key=f"new_item_{campo}")
                        nuevo_item[campo] = valor
                
                if st.button("Agregar Item"):
                    if any(nuevo_item.values()):  # Si al menos un campo tiene valor
                        st.session_state.items_formulario.append(nuevo_item)
                        st.rerun()
            else:
                st.info("Este formulario no requiere items adicionales.")
        
        with tab3:
            st.subheader("Generar Documento Final")
            
            # Verificar campos requeridos
            campos_faltantes = []
            for campo in formulario["campos"]:
                if campo["requerido"] and not st.session_state.datos_formulario.get(campo["id"]):
                    campos_faltantes.append(campo["label"])
            
            if campos_faltantes:
                st.warning(f"Campos requeridos faltantes: {', '.join(campos_faltantes)}")
            else:
                st.success("✅ Todos los campos requeridos están completos")
                
                if st.button("🎯 Generar Documento Word", type="primary"):
                    with st.spinner("Generando documento..."):
                        doc = generar_documento_word(
                            st.session_state.formulario_activo, 
                            st.session_state.datos_formulario
                        )
                        
                        # Convertir a bytes
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        # Botón de descarga
                        st.download_button(
                            label="📥 Descargar Documento",
                            data=doc_buffer.getvalue(),
                            file_name=f"{formulario['nombre']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("¡Documento generado exitosamente!")
            
            # Botón para completar
            if st.button("✅ Completar y Volver al Chat"):
                st.session_state.formulario_activo = None
                st.session_state.datos_formulario = {}
                st.session_state.items_formulario = []
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": f"✅ Formulario '{formulario['nombre']}' completado exitosamente. ¿Necesitas ayuda con algo más?"
                })
                st.rerun()

if __name__ == "__main__":
    main()